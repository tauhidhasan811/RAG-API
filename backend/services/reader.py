import fitz
from docx import Document
import pandas as pd
import pytesseract
from PIL import Image
import os
import platform
import sqlite3
# Automatically set tesseract.exe path for Windows if not in PATH
if platform.system() == "Windows":
    default_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(default_path):
        pytesseract.pytesseract.tesseract_cmd = default_path


def read_text(path):
    full_text =[]
    with open(path, 'r', encoding='utf-8') as text: 
        full_text.append(text.read())
    full_text =" ".join(full_text)
    full_text = full_text.replace('\n', " ")
    return full_text



def read_image(path):
    extracted_text =""
    image = Image.open(path)
    extracted_text = pytesseract.image_to_string(image)
    extracted_text = extracted_text.replace('\n', ' ').strip()
    if not extracted_text:
        return 'nodata'
    return extracted_text

def read_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))  # Format as table-like row

    full_text = " ".join(full_text)
    full_text = full_text.replace('\n', " ")
    return full_text

def read_pdf(pdf_path):
    text = ""
    pdf_file = fitz.open(pdf_path)
    for page in pdf_file:
        text += page.get_text()
        text = text.encode('utf-16', 'surrogatepass').decode('utf-16', 'ignore')
        text = text.replace('\n', ' ')
    return text

def read_csv(path):
    csv_text=''
    csv_file = pd.read_csv(path)
    def row_to_text(row):
        return ", ".join([f"{col}: {row[col]}" for col in row.index])
    csv_text = " ".join([row_to_text(data) for _, data in csv_file.iterrows()])
    return csv_text

def read_sql(db_path):
    all_text=""
    query = "SELECT * FROM products"
    conn = sqlite3.connect(db_path)
    df = pd.read_sql_query(query, conn)
    conn.close()

    def row_to_text(row):
        return ", ".join([f"{col}: {row[col]}" for col in row.index])
    all_text = " ".join([row_to_text(row) for _, row in df.iterrows()])
    return all_text

def read_metaData(path):
    df_from_csv = pd.read_csv(path)
    data_from_csv = df_from_csv.to_dict(orient='records')
    return [data_from_csv]


